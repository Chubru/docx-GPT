import docx
import re
from typing import Callable, Dict


class DocxProcessor:
    def __init__(self, prompt: str, max_message_size: int = 2000):
        self.prompt = prompt
        self.max_message_size = max_message_size

    @staticmethod
    def _extract_text_by_id(input_string, target_id):
        pattern = re.compile(r'<!!id=(\d+)>([^<]+)(<!>)?')
        matches = pattern.finditer(input_string)

        extracted_text = None
        error = False

        for match in matches:
            current_id = int(match.group(1))
            current_text = match.group(2)
            closing_tag = match.group(3)

            if str(current_id) == target_id:
                if closing_tag is None:
                    error = True
                    break
                else:
                    extracted_text = current_text
                    break

        if error:
            print("Error: The closing marker is missing or is later than the next opening one.")
            return "-1"
        elif extracted_text is None:
            print(f"Error: ID {target_id} not found.")
            return "-2"
        else:
            return extracted_text

    def _process_text(self, text_map: Dict[str, str], callback: Callable) -> Dict[str, str]:
        processed_text_map = {}
        keys_list = list(text_map.keys())

        while len(keys_list) != 0:
            sending_id = []
            requests = self.prompt

            while len(keys_list) != 0 and len(requests) + len(text_map[keys_list[0]]) < self.max_message_size:
                requests = requests + '<!!id=' + str(keys_list[0]) + '>' + text_map[keys_list[0]] + '<!>'
                sending_id.append(keys_list[0])
                del keys_list[0]

            processed_text = callback(requests)
            processed_text = processed_text.strip('\n')

            for i in range(len(sending_id)):
                string = self._extract_text_by_id(processed_text, sending_id[i])

                if string == '-1' or string == '-2':
                    processed_text_map[str(sending_id[i])] = callback(self.prompt + text_map[str(sending_id[i])])
                    continue

                processed_text_map[str(sending_id[i])] = string

        return processed_text_map

    def process(self, input_file: str, output_file: str, callback: Callable):
        doc = docx.Document(input_file)
        text_map = {}

        index = 0

        for _, paragraph in enumerate(doc.paragraphs):
            if paragraph.text != '':
                text_map[str(index + 1)] = paragraph.text
                index = index + 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text != '':
                            text_map[str(index + 1)] = paragraph.text
                            index = index + 1

        processed_text_map = self._process_text(text_map, callback)

        index = 0
        for _, paragraph in enumerate(doc.paragraphs):
            if str(index + 1) in processed_text_map:
                paragraph.text = processed_text_map[str(index + 1)]
                index = index + 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if str(index + 1) in processed_text_map:
                            paragraph.text = processed_text_map[str(index + 1)]
                            index = index + 1

        doc.save(output_file)
